"""
AI Video Interviewer — Flask Backend
Replaces HuggingFace/Gradio with Google Gemini API.
Now enhanced with data-driven question selection from CSV/JSON datasets.

Endpoints:
  POST /api/generate-question-pool   — 20 questions for a campaign (dataset + Gemini)
  POST /api/select-questions         — 5 questions from pool based on match score
  POST /api/evaluate-answer-ai       — score one answer
  POST /api/generate-evaluation      — full session evaluation report
  POST /api/parse-resume             — PDF/DOCX → skills, experience, match score
  POST /api/transcribe               — audio → text (Google Speech Recognition)
  POST /api/text-to-speech           — text → MP3 URL
  GET  /api/test-question-selection  — test dataset question selection
  GET  /api/dataset-stats            — dataset statistics
"""

import os
import re
import json
import uuid
import secrets
import datetime
import tempfile
import traceback

from flask import Flask, request, jsonify, send_from_directory
from flask_cors import CORS
from dotenv import load_dotenv
from gtts import gTTS
import google.generativeai as genai
import speech_recognition as sr

load_dotenv()

# ── Gemini ─────────────────────────────────────────────────────────────────
GEMINI_KEY = os.getenv('GEMINI_API_KEY', '')
if GEMINI_KEY:
    genai.configure(api_key=GEMINI_KEY)
    gemini_model = genai.GenerativeModel('gemini-2.0-flash')
else:
    gemini_model = None
    print("⚠️  GEMINI_API_KEY not set — AI features will use fallback stubs.")

# ── Flask ───────────────────────────────────────────────────────────────────
app = Flask(__name__, static_folder='static')
CORS(app, resources={r"/api/*": {"origins": "*"}})

AUDIO_DIR = os.path.join(app.static_folder or 'static', 'audio')
os.makedirs(AUDIO_DIR, exist_ok=True)

# ── In-memory store (fallback when Firebase not configured) ─────────────────
_sessions_store: dict = {}      # session_id → evaluation data
_question_pools: dict = {}      # campaign_id → list of questions
_violation_logs: dict = {}      # session_id → list of violations
_links_store: dict = {}         # link_id → link data (email link tracking)

# ── Email Service ──────────────────────────────────────────────────────────
from email_service import send_email, build_interview_email

# ── Sarvam AI TTS ─────────────────────────────────────────────────────────
SARVAM_API_KEY = os.getenv('SARVAM_API_KEY', '')
SARVAM_TTS_URL = 'https://api.sarvam.ai/text-to-speech'
SARVAM_SPEAKERS = {
    'female': {'en-IN': 'anushka', 'hi-IN': 'anushka', 'te-IN': 'anushka'},
    'male':   {'en-IN': 'abhilash', 'hi-IN': 'abhilash', 'te-IN': 'abhilash'},
}

# ── Dataset-driven question service ─────────────────────────────────────────
from question_service import (
    load_all_questions,
    select_questions_by_role,
    select_questions_for_pool,
    get_dataset_stats,
)
# Pre-load datasets at startup
try:
    _dataset = load_all_questions()
    print(f"📚 Dataset loaded: {len(_dataset)} questions ready")
except Exception as e:
    print(f"⚠️  Dataset load failed: {e} — will use Gemini/fallback only")
    _dataset = []

# ── Helpers ─────────────────────────────────────────────────────────────────

def gemini_generate(prompt: str, fallback: str = '') -> str:
    if not gemini_model:
        return fallback
    try:
        response = gemini_model.generate_content(prompt)
        return response.text
    except Exception as e:
        print(f"Gemini error: {e}")
        return fallback


def extract_json(text: str):
    """Pull first JSON object/array from a Gemini response string."""
    # Try to find a JSON block
    match = re.search(r'```json\s*([\s\S]+?)\s*```', text)
    if match:
        return json.loads(match.group(1))
    match = re.search(r'(\[[\s\S]+\]|\{[\s\S]+\})', text)
    if match:
        return json.loads(match.group(1))
    return None


def resolve_public_frontend_url(explicit_url: str = '') -> str:
    explicit_url = (explicit_url or '').strip().rstrip('/')
    public_url = (os.environ.get('PUBLIC_URL', '') or '').strip().rstrip('/')

    if explicit_url and all(host not in explicit_url for host in ('localhost', '127.0.0.1')):
        return explicit_url

    if public_url:
        return public_url

    forwarded_host = (request.headers.get('X-Forwarded-Host') or request.headers.get('Host') or '').split(',')[0].strip()
    forwarded_proto = (request.headers.get('X-Forwarded-Proto') or request.scheme or 'http').split(',')[0].strip()
    if forwarded_host:
        return f'{forwarded_proto}://{forwarded_host}'

    if explicit_url:
        return explicit_url

    return 'http://localhost:5173'


def build_frontend_app_url(frontend_url: str, route_path: str) -> str:
    base_url = (frontend_url or '').strip().rstrip('/')
    normalized_path = route_path if route_path.startswith('/') else f'/{route_path}'
    return f'{base_url}/#{normalized_path}'


def parse_client_datetime(value: str | None):
    if not value:
        return datetime.datetime.utcnow()
    try:
        parsed = datetime.datetime.fromisoformat(value.replace('Z', '+00:00'))
        if parsed.tzinfo:
            parsed = parsed.astimezone(datetime.timezone.utc).replace(tzinfo=None)
        return parsed
    except Exception:
        return datetime.datetime.utcnow()


def _get_firestore_db():
    try:
        from firebase_service import get_db
        return get_db()
    except Exception:
        return None


def find_link_by_room(room_id: str):
    if not room_id:
        return None

    db = _get_firestore_db()
    if db is not None:
        try:
            docs = db.collection('links').where('roomId', '==', room_id).limit(1).stream()
            for doc in docs:
                data = doc.to_dict() or {}
                data.setdefault('linkId', doc.id)
                return data
        except Exception:
            pass

    for link_id, link_data in _links_store.items():
        if link_data.get('roomId') == room_id:
            return {**link_data, 'linkId': link_data.get('linkId', link_id)}
    return None


def load_campaign_context(campaign_id: str):
    if not campaign_id:
        return None

    db = _get_firestore_db()
    if db is None:
        return None

    try:
        snap = db.collection('campaigns').document(campaign_id).get()
        if snap.exists:
            return {'id': snap.id, **(snap.to_dict() or {})}
    except Exception:
        pass
    return None


def build_session_context(room_id: str, explicit_campaign_id: str = ''):
    link_data = find_link_by_room(room_id) or {}
    campaign_id = explicit_campaign_id or link_data.get('campaignId') or ''
    campaign = load_campaign_context(campaign_id) if campaign_id else None
    return {
        'link': link_data,
        'campaignId': campaign_id,
        'campaign': campaign or {},
        'observerRoomId': campaign_id or room_id,
        'campaignTitle': (campaign or {}).get('title') or link_data.get('jobTitle') or '',
        'hrId': (campaign or {}).get('hrId') or link_data.get('createdBy') or '',
    }


# ── Skill list for resume parsing ───────────────────────────────────────────
TECH_SKILLS = [
    "python","javascript","typescript","java","c++","c#","go","rust","ruby","php","swift","kotlin",
    "react","vue","angular","node.js","express","django","flask","fastapi","spring","laravel",
    "sql","postgresql","mysql","mongodb","redis","elasticsearch","cassandra","dynamodb",
    "aws","azure","gcp","docker","kubernetes","terraform","ansible","jenkins","github actions",
    "git","linux","rest api","graphql","grpc","kafka","rabbitmq","celery",
    "machine learning","deep learning","nlp","computer vision","pandas","numpy","tensorflow","pytorch",
    "ci/cd","agile","scrum","tdd","microservices","system design","distributed systems",
]

def parse_skills(text: str, required: str = '') -> dict:
    text_l = text.lower()
    found   = [s for s in TECH_SKILLS if s in text_l]
    req_list = [r.strip().lower() for r in required.split(',') if r.strip()]
    matched  = [r for r in req_list if any(r in s or s in r for s in found)] if req_list else found[:10]
    years_m  = re.search(r'(\d+)\+?\s*years?', text, re.I)
    years    = int(years_m.group(1)) if years_m else 0
    score    = round(len(matched) / max(len(req_list), 1) * 100) if req_list else min(len(found) * 8, 100)
    return {
        'skills': found,
        'matched_skills': matched,
        'years_experience': years,
        'match_score': min(score, 100),
    }


def select_by_difficulty(pool: list, score: int, n: int = 5) -> list:
    if score >= 70:
        difficulty = 'advanced'
    elif score >= 30:
        difficulty = 'intermediate'
    else:
        difficulty = 'basic'

    preferred = [q for q in pool if q.get('difficulty') == difficulty]
    others    = [q for q in pool if q.get('difficulty') != difficulty]

    preferred.sort(key=lambda q: q.get('usedCount', 0))
    others.sort(key=lambda q: q.get('usedCount', 0))

    selected = (preferred + others)[:n]
    return selected


# ══════════════════════════════════════════════════════════════════════════════
# ENDPOINTS
# ══════════════════════════════════════════════════════════════════════════════

@app.route('/api/generate-question-pool', methods=['POST'])
def generate_question_pool():
    """Generate 20 interview questions — tries dataset first, supplements with Gemini."""
    data = request.json or {}
    job_title    = data.get('job_title', 'Software Engineer')
    job_desc     = data.get('job_description', '')
    skills       = data.get('required_skills', '')
    exp_years    = data.get('experience_years', '0-2')
    use_dataset  = data.get('use_dataset', True)  # allow opting out

    questions = []

    # Step 1: Try dataset-driven selection
    if use_dataset and _dataset:
        try:
            questions = select_questions_for_pool(
                role=job_title,
                description=job_desc,
                skills=skills,
                num_questions=20,
            )
            print(f"[generate-pool] Dataset returned {len(questions)} questions for '{job_title}'")
        except Exception as e:
            print(f"[generate-pool] Dataset selection error: {e}")
            questions = []

    # Step 2: If dataset gave fewer than 20, supplement with Gemini
    if len(questions) < 20 and gemini_model:
        needed = 20 - len(questions)
        prompt = f"""
You are an expert technical interviewer. Generate exactly {needed} interview questions for the role:
Title: {job_title}
Experience: {exp_years} years
Required skills: {skills}
Job description summary: {job_desc[:800]}

Create a mix of basic, intermediate, and advanced questions.

For EACH question provide:
- "text": the question
- "difficulty": "basic" | "intermediate" | "advanced"
- "type": "text" or "code" (mark some as "code" for technical roles)
- "modelAnswer": a concise ideal answer (2-3 sentences)
- "rubric": what to look for when scoring (1 sentence)

Return ONLY a valid JSON array (no markdown, no extra text):
[
  {{"id": "q1", "text": "...", "difficulty": "basic", "type": "text", "modelAnswer": "...", "rubric": "...", "usedCount": 0}},
  ...
]
"""
        raw = gemini_generate(prompt)
        try:
            gemini_qs = extract_json(raw)
            if isinstance(gemini_qs, list):
                questions.extend(gemini_qs)
                print(f"[generate-pool] Gemini supplemented {len(gemini_qs)} questions")
        except Exception:
            pass

    # Step 3: Final fallback to stubs
    if len(questions) < 5:
        questions = _stub_questions(job_title, 20)

    # Ensure IDs and usedCount
    for i, q in enumerate(questions):
        q.setdefault('id', f'q{i+1}')
        q.setdefault('usedCount', 0)
        q.setdefault('type', 'text')

    return jsonify({'questions': questions[:20], 'count': min(len(questions), 20)})


@app.route('/api/select-questions', methods=['POST'])
def select_questions():
    """Return 5 questions — tries dataset first, then Gemini, then Firestore pool."""
    data        = request.json or {}
    campaign_id = data.get('campaign_id', '')
    match_score = int(data.get('match_score', 50))
    pool        = data.get('pool', [])
    job_title   = data.get('job_title', '')
    job_desc    = data.get('job_description', '')

    # Determine difficulty preference from match score
    if match_score >= 70:
        diff_pref = 'advanced'
    elif match_score >= 30:
        diff_pref = 'intermediate'
    else:
        diff_pref = 'basic'

    # Step 1: Try dataset-driven selection if job_title is provided and no pool given
    if job_title and _dataset and not pool:
        try:
            dataset_qs = select_questions_by_role(
                role=job_title,
                description=job_desc,
                num_questions=5,
                difficulty=diff_pref,
            )
            if len(dataset_qs) >= 3:
                pool = dataset_qs
                print(f"[select-questions] Dataset returned {len(pool)} questions for '{job_title}'")
        except Exception as e:
            print(f"[select-questions] Dataset error: {e}")

    # Step 2: Try Gemini if still no pool
    if not pool and job_title and gemini_model:
        prompt = f"""
You are a senior technical interviewer. Generate exactly 5 interview questions STRICTLY for this role:
Role: {job_title}
Job Description: {job_desc[:600] if job_desc else 'Not provided'}
Candidate match score: {match_score}%

Rules:
- ALL questions must be directly relevant to the "{job_title}" role
- Do NOT ask generic questions like "tell me about yourself"
- Questions must test skills, knowledge, and problem-solving specific to this role
- If match_score >= 70: ask mostly advanced questions
- If match_score 30-69: ask intermediate questions
- If match_score < 30: ask basic questions
- Include 1 coding question (type: "code") for technical roles
- Each question must be unique and specific

Return ONLY a valid JSON array:
[
  {{"id": "q1", "text": "...", "difficulty": "basic|intermediate|advanced", "type": "text|code", "modelAnswer": "...", "rubric": "..."}},
  ...
]
"""
        raw = gemini_generate(prompt)
        try:
            generated = extract_json(raw)
            if isinstance(generated, list) and len(generated) >= 3:
                for i, q in enumerate(generated):
                    q.setdefault('id', f'q{i+1}')
                    q.setdefault('type', 'text')
                    q.setdefault('difficulty', 'intermediate')
                pool = generated
        except Exception:
            pass

    # Step 3: Fallback to Firestore pool
    if not pool:
        try:
            from firebase_service import get_db
            db   = get_db()
            snap = db.collection('campaigns').document(campaign_id).get()
            if snap.exists:
                pool = snap.to_dict().get('questionPool', [])
        except Exception:
            pool = _question_pools.get(campaign_id, [])

    # Step 4: Final fallback to stubs
    if not pool:
        pool = _stub_questions(job_title or 'General', 20)

    selected = select_by_difficulty(pool, match_score, n=5)
    return jsonify({'questions': selected, 'difficulty_band': _band(match_score)})


@app.route('/api/evaluate-answer-ai', methods=['POST'])
def evaluate_answer_ai():
    """Score a single answer 1–10 with feedback."""
    data         = request.json or {}
    question     = data.get('question', '')
    answer       = data.get('answer', '')
    model_answer = data.get('model_answer', '')
    rubric       = data.get('rubric', '')
    session_id   = data.get('session_id', '')
    q_index      = data.get('question_index', 0)

    prompt = f"""
You are an expert interviewer. Score the candidate's answer.

Question: {question}
Model Answer: {model_answer}
Rubric: {rubric}
Candidate's Answer: {answer}

Return ONLY valid JSON:
{{"score": <integer 1-10>, "feedback": "<2-3 sentence feedback>", "strengths": "<what was good>", "improvements": "<what to improve>"}}
"""
    raw = gemini_generate(prompt, fallback='{"score": 5, "feedback": "Average response.", "strengths": "", "improvements": ""}')
    try:
        result = extract_json(raw) or json.loads(raw)
    except Exception:
        result = {"score": 5, "feedback": raw[:200], "strengths": "", "improvements": ""}

    # Persist in memory (Firestore update can be added here)
    if session_id:
        if session_id not in _sessions_store:
            _sessions_store[session_id] = {'questionEvals': {}}
        _sessions_store[session_id]['questionEvals'][str(q_index)] = result

    return jsonify(result)


@app.route('/api/generate-evaluation', methods=['POST'])
def generate_evaluation():
    """Generate a full session evaluation report."""
    data       = request.json or {}
    session_id = data.get('session_id', '')
    answers    = data.get('answers', [])
    room_id    = data.get('room_id', '')
    context    = build_session_context(room_id, data.get('campaign_id', ''))
    violations = data.get('violations', []) or _violation_logs.get(session_id, [])

    q_evals = {}
    if session_id and session_id in _sessions_store:
        q_evals = _sessions_store[session_id].get('questionEvals', {})

    scores = [v.get('score', 5) for v in q_evals.values()]
    avg    = round(sum(scores) / len(scores) * 10) if scores else 50

    answers_text = '\n'.join(
        f"Q{i+1}: {a.get('question','')}\nA: {a.get('answer','')[:400]}"
        for i, a in enumerate(answers[:10])
    )

    prompt = f"""
You are a senior technical interviewer writing an evaluation report.

Interview answers:
{answers_text}

Per-question scores: {json.dumps(q_evals)}

Write a concise professional evaluation report. Return ONLY valid JSON:
{{
  "overallScore": <integer 0-100>,
  "summary": "<3-4 sentence overall assessment>",
  "strengths": ["<strength 1>", "<strength 2>", "<strength 3>"],
  "areasToImprove": ["<area 1>", "<area 2>"],
  "recommendation": "hire" | "consider" | "pass",
  "questionEvals": {{
    "0": {{"score": <1-10>, "feedback": "..."}},
    ...
  }}
}}
"""
    raw = gemini_generate(prompt)
    try:
        result = extract_json(raw) or json.loads(raw)
    except Exception:
        result = {
            "overallScore": avg,
            "summary": "Evaluation generated from individual scores.",
            "strengths": [],
            "areasToImprove": [],
            "recommendation": "consider",
            "questionEvals": q_evals,
        }

    result['overallScore'] = result.get('overallScore', avg)

    # Update Firestore session
    if session_id:
        session_doc = {
            'sessionId': session_id,
            'roomId': room_id,
            'observerRoomId': context.get('observerRoomId'),
            'candidateId': data.get('candidate_id', ''),
            'candidateEmail': data.get('candidate_email', ''),
            'candidateName': data.get('candidateName', ''),
            'jobTitle': data.get('jobTitle', ''),
            'campaignId': context.get('campaignId') or data.get('campaign_id', ''),
            'campaignTitle': context.get('campaignTitle'),
            'hrId': context.get('hrId') or data.get('hr_id', ''),
            'matchScore': data.get('match_score'),
            'answers': answers,
            'duration': data.get('duration', 0),
            'violations': violations,
            'totalViolations': len(violations),
            'questionEvals': result.get('questionEvals', q_evals),
            'evaluation': result,
            'overallScore': result['overallScore'],
            'status': 'completed',
            'endedAt': datetime.datetime.utcnow(),
            'startedAt': parse_client_datetime(data.get('started_at')),
        }
        try:
            db = _get_firestore_db()
            if db is None:
                raise RuntimeError('Firestore unavailable')
            db.collection('sessions').document(session_id).set(session_doc, merge=True)
        except Exception:
            _sessions_store.setdefault(session_id, {}).update(session_doc)

    return jsonify(result)


@app.route('/api/interview/start', methods=['POST'])
def start_interview():
    """Create or update a session record when an interview begins."""
    data = request.json or {}
    session_id = data.get('session_id', '').strip()
    room_id = data.get('room_id', '').strip()

    if not session_id or not room_id:
        return jsonify({'error': 'session_id and room_id are required'}), 400

    context = build_session_context(room_id, data.get('campaign_id', ''))
    session_doc = {
        'sessionId': session_id,
        'roomId': room_id,
        'observerRoomId': context.get('observerRoomId'),
        'candidateId': data.get('candidate_id', ''),
        'candidateEmail': data.get('candidate_email', ''),
        'candidateName': data.get('candidate_name', ''),
        'jobTitle': data.get('job_title', '') or context.get('campaignTitle'),
        'campaignId': context.get('campaignId'),
        'campaignTitle': context.get('campaignTitle'),
        'hrId': context.get('hrId'),
        'matchScore': data.get('match_score'),
        'status': 'in-progress',
        'startedAt': parse_client_datetime(data.get('started_at')),
    }

    _sessions_store.setdefault(session_id, {}).update({
        'candidateId': session_doc['candidateId'],
        'candidateName': session_doc['candidateName'],
        'campaignId': session_doc['campaignId'],
        'campaignTitle': session_doc['campaignTitle'],
        'hrId': session_doc['hrId'],
        'matchScore': session_doc['matchScore'],
        'status': session_doc['status'],
        'startedAt': session_doc['startedAt'].isoformat(),
    })

    try:
        db = _get_firestore_db()
        if db is None:
            raise RuntimeError('Firestore unavailable')
        db.collection('sessions').document(session_id).set(session_doc, merge=True)
    except Exception as e:
        print(f"[Session Start] Firestore save failed (in-memory only): {e}")

    return jsonify({
        'ok': True,
        'campaignId': context.get('campaignId') or None,
        'campaignTitle': context.get('campaignTitle') or None,
        'hrId': context.get('hrId') or None,
        'observerRoomId': context.get('observerRoomId') or room_id,
    })


@app.route('/api/parse-resume', methods=['POST'])
def parse_resume():
    """Accept PDF or DOCX, return skills, experience years, and match score."""
    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file provided'}), 400

    file         = request.files['resume']
    job_desc     = request.form.get('job_description', '')
    req_skills   = request.form.get('required_skills', '')
    filename     = file.filename.lower()

    text = ''
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        if filename.endswith('.pdf'):
            import pdfplumber
            with pdfplumber.open(tmp_path) as pdf:
                text = '\n'.join(p.extract_text() or '' for p in pdf.pages)
        elif filename.endswith('.docx'):
            from docx import Document
            doc = Document(tmp_path)
            text = '\n'.join(p.text for p in doc.paragraphs)
        else:
            text = file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        return jsonify({'error': f'Could not parse file: {e}'}), 500
    finally:
        if tmp_path:
            try:
                os.remove(tmp_path)
            except Exception:
                pass

    parsed = parse_skills(text, req_skills or job_desc)
    return jsonify(parsed)


@app.route('/api/transcribe', methods=['POST'])
def transcribe_audio():
    """Convert uploaded audio (webm/wav) to text via Google Speech Recognition."""
    if 'audio' not in request.files:
        return jsonify({'error': 'No audio file'}), 400

    audio_file = request.files['audio']
    input_path = tempfile.mktemp(suffix='.webm')
    wav_path   = tempfile.mktemp(suffix='.wav')

    try:
        audio_file.save(input_path)

        # Convert to WAV using pydub (requires ffmpeg)
        try:
            from pydub import AudioSegment
            AudioSegment.from_file(input_path).export(wav_path, format='wav')
        except Exception:
            # If pydub/ffmpeg not available, try directly
            wav_path = input_path

        recognizer = sr.Recognizer()
        with sr.AudioFile(wav_path) as source:
            audio_data = recognizer.record(source)
            text = recognizer.recognize_google(audio_data)
    except sr.UnknownValueError:
        text = ''
    except Exception as e:
        return jsonify({'error': str(e), 'text': ''}), 200
    finally:
        for p in [input_path, wav_path]:
            try:
                os.remove(p)
            except Exception:
                pass

    return jsonify({'text': text})


@app.route('/api/text-to-speech', methods=['POST'])
def text_to_speech():
    """Convert text to MP3 and return a URL. Supports multilingual via Sarvam AI."""
    data = request.json or {}
    text = data.get('text', '').strip()
    language = data.get('language', 'en-IN')
    gender = data.get('gender', 'female')
    if not text:
        return jsonify({'error': 'No text provided'}), 400

    text = re.sub(r'\*+', '', text)[:500]   # clean markdown, cap length
    ts   = datetime.datetime.now().strftime('%Y%m%d_%H%M%S_%f')
    fname = f'tts_{ts}.mp3'
    fpath = os.path.join(AUDIO_DIR, fname)

    # Try Sarvam AI first for all languages
    if SARVAM_API_KEY:
        try:
            import requests as req
            speaker = SARVAM_SPEAKERS.get(gender, SARVAM_SPEAKERS['female']).get(language, 'anushka')
            resp = req.post(SARVAM_TTS_URL, headers={
                'api-subscription-key': SARVAM_API_KEY,
                'Content-Type': 'application/json',
            }, json={
                'inputs': [text],
                'target_language_code': language,
                'speaker': speaker,
                'pitch': 0,
                'pace': 1.0,
                'loudness': 1.0,
                'speech_sample_rate': 24000,
            }, timeout=15)
            if resp.status_code == 200:
                resp_data = resp.json()
                # Sarvam returns base64 audio
                audios = resp_data.get('audios', [])
                if audios:
                    import base64
                    audio_bytes = base64.b64decode(audios[0])
                    with open(fpath, 'wb') as f:
                        f.write(audio_bytes)
                    return jsonify({'audio_url': f'/static/audio/{fname}', 'engine': 'sarvam'})
        except Exception as e:
            print(f"[TTS] Sarvam AI failed: {e}")

    # Fallback: gTTS (English only)
    try:
        gtts_lang = 'hi' if language.startswith('hi') else 'te' if language.startswith('te') else 'en'
        gTTS(text=text, lang=gtts_lang, slow=False).save(fpath)
        return jsonify({'audio_url': f'/static/audio/{fname}', 'engine': 'gtts'})
    except Exception as e:
        return jsonify({'error': str(e), 'fallback': 'browser'}), 200


@app.route('/static/audio/<path:filename>')
def serve_audio(filename):
    return send_from_directory(AUDIO_DIR, filename)


@app.route('/api/ats/score', methods=['POST'])
def ats_score():
    """ATS scoring endpoint — receives resume file + job description, returns score and feedback."""
    if 'resume' not in request.files:
        return jsonify({'error': 'No resume file provided'}), 400

    file         = request.files['resume']
    job_desc     = request.form.get('job_description', '')
    req_skills   = request.form.get('required_skills', '')
    filename     = file.filename.lower()

    text = ''
    tmp_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as tmp:
            file.save(tmp.name)
            tmp_path = tmp.name

        if filename.endswith('.pdf'):
            import pdfplumber
            with pdfplumber.open(tmp_path) as pdf:
                text = '\n'.join(p.extract_text() or '' for p in pdf.pages)
        elif filename.endswith('.docx'):
            from docx import Document
            doc = Document(tmp_path)
            text = '\n'.join(p.text for p in doc.paragraphs)
        else:
            text = file.read().decode('utf-8', errors='ignore')
    except Exception as e:
        return jsonify({'error': f'Could not parse file: {e}'}), 500
    finally:
        if tmp_path:
            try: os.remove(tmp_path)
            except: pass

    parsed = parse_skills(text, req_skills or job_desc)

    # Generate AI feedback if Gemini available
    feedback = ''
    if gemini_model and parsed['match_score'] > 0:
        try:
            prompt = f"""Given a candidate with skills: {', '.join(parsed['skills'][:15])} and {parsed['years_experience']} years experience.
Job requires: {req_skills or job_desc[:300]}
Match score: {parsed['match_score']}%.
Give 2-3 sentences of constructive feedback about this candidate's fit. Be specific."""
            feedback = gemini_generate(prompt, '')
        except: pass

    result = {
        **parsed,
        'eligible': parsed['match_score'] >= 60,
        'threshold': 60,
        'feedback': feedback,
        'recommendation': 'eligible' if parsed['match_score'] >= 60 else 'below-threshold',
    }
    return jsonify(result)


@app.route('/api/auto-generate-campaign', methods=['POST'])
def auto_generate_campaign():
    """Auto-generate a complete test campaign with job title, description, skills, and questions."""
    data = request.json or {}
    role_type = data.get('role_type', 'fullstack')  # frontend, backend, fullstack, data, devops, mobile, qa

    role_templates = {
        'frontend': {'title': 'Frontend Developer', 'skills': 'React, JavaScript, TypeScript, CSS, HTML, Responsive Design, Redux, Webpack', 'desc': 'Build modern, responsive web interfaces with React and TypeScript.'},
        'backend': {'title': 'Backend Developer', 'skills': 'Node.js, Python, REST APIs, PostgreSQL, Redis, Docker, Microservices', 'desc': 'Design and build scalable backend services and APIs.'},
        'fullstack': {'title': 'Full Stack Developer', 'skills': 'React, Node.js, TypeScript, PostgreSQL, Docker, REST APIs, Git', 'desc': 'Develop end-to-end web applications with modern tech stack.'},
        'data': {'title': 'Data Engineer', 'skills': 'Python, SQL, Spark, Airflow, AWS, ETL, Data Modeling, Kafka', 'desc': 'Build and maintain data pipelines and infrastructure.'},
        'devops': {'title': 'DevOps Engineer', 'skills': 'AWS, Docker, Kubernetes, Terraform, CI/CD, Linux, Monitoring, Jenkins', 'desc': 'Manage cloud infrastructure and deployment pipelines.'},
        'mobile': {'title': 'Mobile Developer', 'skills': 'React Native, Swift, Kotlin, Flutter, REST APIs, Firebase, App Store', 'desc': 'Build cross-platform mobile applications.'},
        'qa': {'title': 'QA Engineer', 'skills': 'Selenium, Cypress, Jest, API Testing, Performance Testing, CI/CD, Python', 'desc': 'Ensure software quality through automated and manual testing.'},
        'ml': {'title': 'ML Engineer', 'skills': 'Python, TensorFlow, PyTorch, NLP, Computer Vision, MLOps, SQL, Docker', 'desc': 'Build and deploy machine learning models in production.'},
    }

    template = role_templates.get(role_type, role_templates['fullstack'])
    custom_title = data.get('title', template['title'])
    custom_skills = data.get('skills', template['skills'])
    custom_desc = data.get('description', template['desc'])
    exp_level = data.get('experience_level', '3-5')

    # Step 1: Try dataset-driven selection
    questions = []
    if _dataset:
        try:
            questions = select_questions_for_pool(
                role=custom_title,
                description=custom_desc,
                skills=custom_skills,
                num_questions=20,
            )
            print(f"[auto-generate] Dataset returned {len(questions)} questions for '{custom_title}'")
        except Exception as e:
            print(f"[auto-generate] Dataset error: {e}")

    # Step 2: Supplement with Gemini if needed
    if len(questions) < 20 and gemini_model:
        needed = 20 - len(questions)
        prompt = f"""
You are an expert technical interviewer. Generate exactly {needed} interview questions for:
Title: {custom_title}
Experience: {exp_level} years
Skills: {custom_skills}
Description: {custom_desc}

Create a balanced mix of basic, intermediate, and advanced questions.
Include some coding questions (type: "code").

Return ONLY a valid JSON array:
[
  {{"id": "q1", "text": "...", "difficulty": "basic|intermediate|advanced", "type": "text|code", "modelAnswer": "...", "rubric": "...", "usedCount": 0}},
  ...
]
"""
        raw = gemini_generate(prompt)
        try:
            gemini_qs = extract_json(raw)
            if isinstance(gemini_qs, list):
                questions.extend(gemini_qs)
        except Exception:
            pass

    # Step 3: Final fallback
    if len(questions) < 5:
        questions = _stub_questions(custom_title, 20)

    for i, q in enumerate(questions):
        q.setdefault('id', f'q{i+1}')
        q.setdefault('usedCount', 0)
        q.setdefault('type', 'text')

    campaign = {
        'title': custom_title,
        'jobDescription': custom_desc,
        'requiredSkills': custom_skills,
        'experienceYears': exp_level,
        'questions': questions,
        'questionCount': len(questions),
        'roleType': role_type,
    }

    return jsonify(campaign)


@app.route('/api/candidate/history', methods=['GET'])
def candidate_history():
    """Return interview history for a candidate from Firestore or in-memory store."""
    candidate_id = request.args.get('candidate_id', '')
    if not candidate_id:
        return jsonify({'error': 'candidate_id required'}), 400

    sessions = []
    try:
        from firebase_service import get_db
        db = get_db()
        docs = db.collection('sessions').where('candidateId', '==', candidate_id).order_by('startedAt', direction='DESCENDING').limit(50).stream()
        for doc in docs:
            d = doc.to_dict()
            d['id'] = doc.id
            sessions.append(d)
    except Exception:
        # Fallback: search in-memory store
        for sid, sdata in _sessions_store.items():
            if sdata.get('candidateId') == candidate_id:
                sessions.append({**sdata, 'id': sid})

    return jsonify({'sessions': sessions, 'count': len(sessions)})


@app.route('/api/candidate/analytics', methods=['GET'])
def candidate_analytics():
    """Return analytics summary for a candidate."""
    candidate_id = request.args.get('candidate_id', '')
    if not candidate_id:
        return jsonify({'error': 'candidate_id required'}), 400

    sessions = []
    try:
        from firebase_service import get_db
        db = get_db()
        docs = db.collection('sessions').where('candidateId', '==', candidate_id).stream()
        for doc in docs:
            sessions.append(doc.to_dict())
    except Exception:
        for sid, sdata in _sessions_store.items():
            if sdata.get('candidateId') == candidate_id:
                sessions.append(sdata)

    total = len(sessions)
    completed = [s for s in sessions if s.get('status') == 'completed']
    scores = [s.get('overallScore', 0) for s in completed if s.get('overallScore')]
    avg_score = round(sum(scores) / len(scores)) if scores else 0

    return jsonify({
        'totalInterviews': total,
        'completed': len(completed),
        'inProgress': total - len(completed),
        'averageScore': avg_score,
        'highestScore': max(scores) if scores else 0,
        'lowestScore': min(scores) if scores else 0,
        'scores': scores,
    })


@app.route('/api/hr/candidates', methods=['GET'])
def hr_candidates():
    """Return all candidates from Firestore (HR view)."""
    candidates = []
    try:
        from firebase_service import get_db
        db = get_db()
        docs = db.collection('users').where('role', '==', 'candidate').limit(200).stream()
        for doc in docs:
            d = doc.to_dict()
            d['id'] = doc.id
            # Remove sensitive fields
            d.pop('password', None)
            candidates.append(d)
    except Exception:
        pass

    return jsonify({'candidates': candidates, 'count': len(candidates)})


@app.route('/api/hr/analytics', methods=['GET'])
def hr_analytics():
    """Return analytics summary for HR dashboard."""
    sessions = []
    campaigns = []
    try:
        from firebase_service import get_db
        db = get_db()
        # Get all sessions
        for doc in db.collection('sessions').limit(500).stream():
            sessions.append(doc.to_dict())
        # Get campaigns
        for doc in db.collection('campaigns').limit(100).stream():
            campaigns.append(doc.to_dict())
    except Exception:
        # Fallback to in-memory
        sessions = list(_sessions_store.values())

    completed = [s for s in sessions if s.get('status') == 'completed']
    scores = [s.get('overallScore', 0) for s in completed if s.get('overallScore')]
    avg_score = round(sum(scores) / len(scores)) if scores else 0

    # Recommendation breakdown
    recs = {}
    for s in completed:
        rec = s.get('evaluation', {}).get('recommendation', 'pending')
        recs[rec] = recs.get(rec, 0) + 1

    # Score distribution
    dist = {'0-20': 0, '21-40': 0, '41-60': 0, '61-80': 0, '81-100': 0}
    for sc in scores:
        if sc <= 20: dist['0-20'] += 1
        elif sc <= 40: dist['21-40'] += 1
        elif sc <= 60: dist['41-60'] += 1
        elif sc <= 80: dist['61-80'] += 1
        else: dist['81-100'] += 1

    return jsonify({
        'totalInterviews': len(sessions),
        'completedInterviews': len(completed),
        'activeCampaigns': len(campaigns),
        'averageScore': avg_score,
        'passRate': round(len([s for s in scores if s >= 60]) / max(len(scores), 1) * 100),
        'recommendations': recs,
        'scoreDistribution': dist,
        'scores': scores[-30:],  # Last 30 for chart
    })


@app.route('/api/test-question-selection', methods=['GET'])
def test_question_selection():
    """Test endpoint — select questions from dataset by role and description."""
    role = request.args.get('role', 'Software Engineer')
    description = request.args.get('description', '')
    difficulty = request.args.get('difficulty', '')
    num = min(int(request.args.get('num', 5)), 20)

    if not _dataset:
        return jsonify({'error': 'Dataset not loaded', 'questions': []}), 200

    questions = select_questions_by_role(
        role=role,
        description=description,
        num_questions=num,
        difficulty=difficulty,
    )

    return jsonify({
        'role': role,
        'description': description,
        'difficulty': difficulty or 'auto',
        'count': len(questions),
        'questions': questions,
    })


@app.route('/api/dataset-stats', methods=['GET'])
def dataset_stats():
    """Return statistics about the loaded question dataset."""
    stats = get_dataset_stats()
    return jsonify(stats)


# ── Violation Logging ──────────────────────────────────────────────────────

@app.route('/api/interview/violation', methods=['POST'])
def log_violation():
    """Log a proctoring violation for a session."""
    data = request.json or {}
    session_id = data.get('session_id', '')
    violation = {
        'type': data.get('type', 'unknown'),
        'timestamp': data.get('timestamp', datetime.datetime.now().isoformat()),
        'details': data.get('details', ''),
    }

    if not session_id:
        return jsonify({'error': 'session_id required'}), 400

    _violation_logs.setdefault(session_id, []).append(violation)

    # Also try Firestore
    try:
        from firebase_service import get_db
        fdb = get_db()
        from google.cloud.firestore_v1 import ArrayUnion
        fdb.collection('sessions').document(session_id).update({
            'violations': ArrayUnion([violation])
        })
    except Exception as e:
        print(f"[Violation] Firestore update failed (in-memory only): {e}")

    return jsonify({'ok': True, 'count': len(_violation_logs[session_id])})


@app.route('/api/interview/report/<session_id>', methods=['GET'])
def get_interview_report(session_id):
    """Get comprehensive interview report for a session."""
    report = {}

    # Try Firestore first
    try:
        from firebase_service import get_db
        db = get_db()
        snap = db.collection('sessions').document(session_id).get()
        if snap.exists:
            report = snap.to_dict()
    except Exception:
        pass

    # Merge in-memory data
    if session_id in _sessions_store:
        mem = _sessions_store[session_id]
        if not report.get('evaluation'):
            report['evaluation'] = mem
        if not report.get('questionEvals'):
            report['questionEvals'] = mem.get('questionEvals', {})

    # Add violations
    violations = report.get('violations', []) or _violation_logs.get(session_id, [])
    report['violations'] = violations

    # Compute violation summary
    violation_summary = {}
    for v in violations:
        vtype = v.get('type', 'unknown')
        violation_summary[vtype] = violation_summary.get(vtype, 0) + 1
    report['violationSummary'] = violation_summary
    report['totalViolations'] = len(violations)

    # Compute recommendation from evaluation
    eval_data = report.get('evaluation', {})
    score = eval_data.get('overallScore', report.get('overallScore', 0))
    if score >= 70:
        recommendation = 'hire'
    elif score >= 50:
        recommendation = 'maybe'
    else:
        recommendation = 'reject'
    report['recommendation'] = eval_data.get('recommendation', recommendation)
    report['overallScore'] = score

    # Check if we actually found any real data (not just our computed fields)
    has_data = (
        session_id in _sessions_store
        or session_id in _violation_logs
        or report.get('evaluation')
        or report.get('answers')
        or report.get('candidateName')
    )
    if not has_data:
        return jsonify({'error': 'Session not found'}), 404

    return jsonify(report)


@app.route('/api/health')
def health():
    return jsonify({
        'status': 'ok',
        'gemini': bool(GEMINI_KEY),
        'dataset_loaded': len(_dataset) > 0,
        'dataset_size': len(_dataset),
    })


# ── Stub helpers ─────────────────────────────────────────────────────────────

def _band(score):
    if score >= 70: return 'advanced'
    if score >= 30: return 'intermediate'
    return 'basic'

def _stub_questions(role, n):
    diffs = ['basic', 'basic', 'intermediate', 'intermediate', 'advanced']
    return [
        {
            'id': f'q{i+1}',
            'text': f'[Fallback Q{i+1}] Describe your experience with {role} and related technologies.',
            'difficulty': diffs[i % len(diffs)],
            'type': 'text',
            'modelAnswer': 'A strong answer covers practical experience, specific examples, and measurable outcomes.',
            'rubric': 'Look for specific examples and depth of knowledge.',
            'usedCount': 0,
        }
        for i in range(n)
    ]


# ═══════════════════════════════════════════════════════════════════════════
# LINK GENERATION & EMAIL
# ═══════════════════════════════════════════════════════════════════════════

@app.route('/api/generate-link', methods=['POST'])
def generate_link():
    """
    Generate a unique interview/mock link and optionally send it by email.

    Body JSON:
      type         — "interview" | "mock"
      role         — "hr" | "candidate" (who is creating)
      userId       — UID of the creator
      email        — candidate email (required for HR, optional for candidate)
      candidateName — name of candidate (optional)
      jobTitle     — job title (optional)
      campaignId   — campaign to link to (optional)
      note         — HR note to include in email (optional)

    Returns:
      { success, link, linkId, emailSent }
    """
    try:
        data = request.json or {}
        link_type = data.get('type', 'interview')
        role = data.get('role', 'candidate')
        user_id = data.get('userId', '')
        email = data.get('email', '').strip()
        candidate_name = data.get('candidateName', '').strip()
        job_title = data.get('jobTitle', '').strip()
        campaign_id = data.get('campaignId', '')
        note = data.get('note', '').strip()
        frontend_url = resolve_public_frontend_url(data.get('frontendUrl', 'http://localhost:5173'))

        # ── Validation ──
        if role == 'hr' and link_type == 'interview' and not email:
            return jsonify({"error": "Candidate email is required for interview links"}), 400

        # ── Generate unique link ID ──
        link_id = secrets.token_urlsafe(16)

        # ── Build the full URL ──
        if link_type == 'mock':
            full_link = build_frontend_app_url(frontend_url, f"/mock?token={link_id}")
        else:
            room_id = campaign_id or f"link-{link_id[:12]}-{secrets.token_hex(4)}"
            full_link = build_frontend_app_url(frontend_url, f"/interview/{room_id}")

        # ── Store link data ──
        link_data = {
            "linkId": link_id,
            "type": link_type,
            "createdBy": user_id,
            "forEmail": email,
            "candidateName": candidate_name,
            "jobTitle": job_title,
            "campaignId": campaign_id,
            "roomId": room_id if link_type != 'mock' else None,
            "fullLink": full_link,
            "createdAt": datetime.datetime.utcnow().isoformat(),
            "used": False,
            "note": note,
        }

        # Try Firestore first
        firestore_saved = False
        try:
            from firebase_service import get_db
            fdb = get_db()
            fdb.collection('links').document(link_id).set(link_data)
            firestore_saved = True
        except Exception:
            pass

        # Always save to in-memory store
        _links_store[link_id] = link_data

        # ── Send email via Resend ──
        # NOTE: Resend free tier can only send to the verified owner email.
        # We always send to the REPLY_TO_EMAIL and include candidate info in the body.
        email_result = {"sent": False, "method": "none"}
        email_error = None
        verified_email = os.environ.get('REPLY_TO_EMAIL', 'giddisathwik.ai@gmail.com')
        if email:
            is_mock = link_type == 'mock'
            html_body = build_interview_email(
                candidate_name=candidate_name,
                link=full_link,
                job_title=job_title,
                note=note,
                is_mock=is_mock,
            )
            # Add forwarding note if sending to a different candidate
            if email.lower() != verified_email.lower():
                html_body = html_body.replace(
                    '</body>',
                    f'<p style="margin-top:20px;padding:12px;background:#f5f3ff;border:1px solid #e8e4f0;font-size:13px;color:#555;">'
                    f'<strong>Forward this email to:</strong> {email}<br>'
                    f'<em>Resend free tier only delivers to the verified account. Please forward this to the candidate.</em>'
                    f'</p></body>'
                )
            subject = f"{'Mock ' if is_mock else ''}Interview Invitation{' — ' + job_title if job_title else ''}"
            if email.lower() != verified_email.lower():
                subject = f"[Forward to {email}] {subject}"
            try:
                # Always send to verified email (free tier limitation)
                email_result = send_email(verified_email, subject, html_body)
            except (ValueError, RuntimeError) as email_err:
                email_error = str(email_err)
                print(f"[LINK] Link {link_id} created but email failed: {email_error}")

        return jsonify({
            "success": True,
            "link": full_link,
            "linkId": link_id,
            "roomId": link_data.get("roomId"),
            "emailSent": email_result.get("sent", False),
            "emailMethod": email_result.get("method", "none"),
            "emailError": email_error,
            "emailDeliveredTo": verified_email if email_result.get("sent", False) else None,
            "emailRequestedFor": email or None,
            "emailForwardRequired": bool(email and email.lower() != verified_email.lower()),
            "storedIn": "firestore" if firestore_saved else "memory",
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/api/links', methods=['GET'])
def get_links():
    """
    Get all links created by a specific user.

    Query params:
      userId — UID of the creator
      role   — "hr" | "candidate" (optional, for filtering)

    Returns:
      { links: [...] }
    """
    try:
        user_id = request.args.get('userId', '')
        if not user_id:
            return jsonify({"error": "userId is required"}), 400

        links = []

        # Try Firestore first
        try:
            from firebase_service import get_db
            fdb = get_db()
            docs = fdb.collection('links').where('createdBy', '==', user_id) \
                       .order_by('createdAt', direction='DESCENDING').stream()
            for doc in docs:
                d = doc.to_dict()
                links.append(d)
        except Exception:
            pass

        # Merge in-memory links (avoid duplicates)
        seen_ids = {l['linkId'] for l in links}
        for lid, ldata in _links_store.items():
            if ldata.get('createdBy') == user_id and lid not in seen_ids:
                links.append(ldata)

        # Sort by creation date, newest first
        links.sort(key=lambda x: x.get('createdAt', ''), reverse=True)

        return jsonify({"links": links})

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


@app.route('/api/link/<link_id>', methods=['GET'])
def get_link(link_id):
    """
    Validate and return a link by its ID.
    Used when someone clicks a link from an email.

    Returns:
      { found, link: {...} }
    """
    try:
        link_data = None

        # Try Firestore first
        try:
            from firebase_service import get_db
            fdb = get_db()
            doc = fdb.collection('links').document(link_id).get()
            if doc.exists:
                link_data = doc.to_dict()
        except Exception:
            pass

        # Fallback to in-memory
        if not link_data:
            link_data = _links_store.get(link_id)

        if not link_data:
            return jsonify({"found": False, "error": "Link not found"}), 404

        # Mark as used
        link_data['used'] = True
        try:
            from firebase_service import get_db
            fdb = get_db()
            fdb.collection('links').document(link_id).update({"used": True})
        except Exception:
            pass
        _links_store[link_id] = link_data

        return jsonify({
            "found": True,
            "link": {
                "linkId": link_data.get("linkId"),
                "type": link_data.get("type"),
                "roomId": link_data.get("roomId"),
                "fullLink": link_data.get("fullLink"),
                "jobTitle": link_data.get("jobTitle"),
                "campaignId": link_data.get("campaignId"),
                "used": True,
            }
        })

    except Exception as e:
        traceback.print_exc()
        return jsonify({"error": str(e)}), 500


if __name__ == '__main__':
    port = int(os.environ.get('PORT', '5001'))
    debug = os.environ.get('FLASK_DEBUG', '1') == '1'
    print(f"🚀 Starting AI Interviewer Flask backend on port {port}…")
    print(f"   Gemini AI: {'✅ enabled' if GEMINI_KEY else '⚠️  disabled (set GEMINI_API_KEY)'}")
    app.run(debug=debug, host='0.0.0.0', port=port)
